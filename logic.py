import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import matplotlib.font_manager as fm
import io
import os
import requests
import streamlit as st

# -----------------------------------------------------------------------------
# 1. AI 모델 설정
# -----------------------------------------------------------------------------
model = None
api_key_status = "키 없음"

try:
    # secrets.toml에서 google_api_key를 찾습니다.
    # [google_api_key] 섹션이 없어도, top-level이나 다른 방식으로 로드될 수 있음
    if "google_api_key" in st.secrets:
        api_key = st.secrets["google_api_key"]
        # 혹시 딕셔너리 형태라면 그 안에서 찾기
        if isinstance(api_key, dict) and "api_key" in api_key:
            api_key = api_key["api_key"]
            
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('models/gemini-2.5-flash')
        api_key_status = "설정 완료"
    else:
        api_key_status = "Secrets에 google_api_key 없음"
except Exception as e:
    api_key_status = f"설정 오류: {e}"
    print(f"모델 설정 오류: {e}")

# -----------------------------------------------------------------------------
# 2. 폰트 자동 다운로드
# -----------------------------------------------------------------------------
FONT_FILE = "NanumGothic.ttf"
FONT_URL = "https://github.com/google/fonts/raw/main/ofl/nanumgothic/NanumGothic-Regular.ttf"

def get_korean_font():
    if not os.path.exists(FONT_FILE):
        try:
            response = requests.get(FONT_URL)
            with open(FONT_FILE, "wb") as f: f.write(response.content)
        except: pass
    try: return fm.FontProperties(fname=FONT_FILE)
    except: return fm.FontProperties(family="sans-serif")

# -----------------------------------------------------------------------------
# 3. 문서 유틸리티
# -----------------------------------------------------------------------------
def set_read_only(doc):
    settings = doc.settings.element
    protection = OxmlElement('w:documentProtection')
    protection.set(qn('w:edit'), 'readOnly')
    protection.set(qn('w:enforcement'), '1')
    settings.append(protection)

def set_font(run, font_name='맑은 고딕', font_size=11, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color: run.font.color.rgb = color

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

# -----------------------------------------------------------------------------
# 4. 그래프 생성
# -----------------------------------------------------------------------------
def create_plot_image(code_snippet):
    kor_font = get_korean_font()
    plt.clf()
    plt.style.use('default')
    plt.rcParams['axes.unicode_minus'] = False
    
    fig, ax = plt.subplots(figsize=(5, 4))
    
    try:
        local_scope = {'plt': plt, 'ax': ax, 'fig': fig, 'patches': patches}
        exec(code_snippet, {}, local_scope)
        
        ax.autoscale(enable=True, axis='both', tight=True)
        ax.set_aspect('equal', adjustable='box')
        
        if ax.axison:
            if kor_font:
                for item in ([ax.title, ax.xaxis.label, ax.yaxis.label] + ax.get_xticklabels() + ax.get_yticklabels()):
                    item.set_fontproperties(kor_font)
                for child in ax.get_children():
                    if isinstance(child, plt.Text): child.set_fontproperties(kor_font)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
        else:
            if kor_font:
                for child in ax.get_children():
                    if isinstance(child, plt.Text): child.set_fontproperties(kor_font)

        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception as e:
        plt.close(fig)
        return None

# -----------------------------------------------------------------------------
# 5. 콘텐츠 생성 (에러 처리 강화)
# -----------------------------------------------------------------------------
def create_error_docx(error_msg):
    doc = Document()
    doc.add_heading('⚠️ 문제 생성 실패', 0)
    p = doc.add_paragraph()
    run = p.add_run(f"오류 내용: {error_msg}")
    run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph("\n[해결 방법]")
    doc.add_paragraph("1. secrets.toml 파일에 'google_api_key'가 있는지 확인하세요.")
    doc.add_paragraph("2. API Key가 올바른지(오타, 공백) 확인하세요.")
    doc.add_paragraph(f"현재 상태: {api_key_status}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_math_docx(school, grade, topic, difficulty, count, is_commercial=False):
    
    if not model:
        return create_error_docx("AI 모델(Gemini)이 설정되지 않았습니다. API Key를 확인해주세요.")

    system_prompt = f"""
    당신은 대한민국 수학 최상위권 교재 집필진입니다.
    요청: {school} {grade}학년 '{topic}' (난이도: {difficulty}) {count}문제.

    [작성 규칙]
    1. 사고력, 문장제, 도형 위주 출제.
    2. 모든 문제에 Python Matplotlib 시각화 코드 필수.
    3. 그림은 '교과서 삽화' 스타일 (축 숨기기, patches 사용).
    
    [출력 형식]
    문제 1: ...
    CODE_START
    ...
    CODE_END
    정답: ...
    @@@
    """

    try:
        response = model.generate_content(system_prompt)
        raw_text = response.text
    except Exception as e:
        return create_error_docx(f"AI 응답 오류: {str(e)}")
    
    # 여기서부터는 정상 생성 로직
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5); section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5); section.bottom_margin = Inches(0.5)
    
    # 헤더
    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = 'Table Grid'
    cell_left = header_table.cell(0, 0)
    p_logo = cell_left.paragraphs[0]
    try: p_logo.add_run().add_picture('logo.png', height=Inches(0.6))
    except: 
        run = p_logo.add_run("지니매쓰")
        set_font(run, font_size=16, bold=True, color=RGBColor(0, 51, 153))
    
    cell_right = header_table.cell(0, 1)
    p_info = cell_right.paragraphs[0]
    run_info = p_info.add_run(f"{topic} ({difficulty})  |  지니매쓰")
    set_font(run_info, font_size=11)
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    for row in header_table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right']:
                el = OxmlElement(f'w:{border}')
                el.set(qn('w:val'), 'nil')
                tcBorders.append(el)
            tcPr.append(tcBorders)

    doc.add_paragraph("") 
    
    items = raw_text.split('@@@')
    answers_list = []
    page_prob_count = 0
    
    for idx, item in enumerate(items):
        if idx >= count: break
        if not item.strip(): continue
        
        lines = item.strip().split('\n')
        mode = "TEXT"
        temp_q, code_text, answer_text = [], "", ""
        
        for line in lines:
            if "CODE_START" in line: mode = "CODE"
            elif "CODE_END" in line: mode = "TEXT"
            elif line.startswith("정답:"): answer_text = line.replace("정답:", "").strip()
            else:
                if mode == "CODE": code_text += line + "\n"
                elif mode == "TEXT" and not line.startswith("문제"): temp_q.append(line)
        
        q_text = "\n".join(temp_q).strip()
        answers_list.append(f"{idx+1}. {answer_text}")
        
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(4.8)
        table.columns[1].width = Inches(2.5)
        
        cell_q = table.cell(0, 0)
        p_num = cell_q.paragraphs[0]
        run_num = p_num.add_run(f"{idx+1}. ")
        set_font(run_num, font_size=13, bold=True)
        
        run_q = cell_q.add_paragraph().add_run(q_text)
        set_font(run_q, font_size=11)
        
        if code_text.strip():
            img_buf = create_plot_image(code_text)
            if img_buf:
                p_img = cell_q.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(img_buf, width=Inches(3.5))

        cell_a = table.cell(0, 1)
        p_sol = cell_a.paragraphs[0]
        run_sol = p_sol.add_run("[ 풀 이 ]")
        set_font(run_sol, font_size=10, color=RGBColor(150, 150, 150))
        p_sol.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        tcPr = cell_a._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '6') 
        left.set(qn('w:color'), 'E0E0E0')
        tcBorders.append(left)
        tcPr.append(tcBorders)

        doc.add_paragraph("")
        
        page_prob_count += 1
        if page_prob_count % 4 == 0:
            if idx < count - 1: doc.add_page_break()
        else:
            p_line = doc.add_paragraph()
            run_line = p_line.add_run("-" * 90)
            set_font(run_line, font_size=8, color=RGBColor(200, 200, 200))
            p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    p_ans = doc.add_paragraph("< 정 답 및 풀 이 >")
    p_ans.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p_ans.runs[0], font_size=16, bold=True)
    doc.add_paragraph("")
    
    ans_table = doc.add_table(rows=(len(answers_list)+1)//2, cols=2)
    ans_table.style = 'Table Grid'
    
    for i, ans in enumerate(answers_list):
        r, c = divmod(i, 2)
        cell = ans_table.cell(r, c)
        cell.text = ans
        for p in cell.paragraphs:
            for r in p.runs: set_font(r, font_size=10)

    section = doc.sections[0]
    footer = section.footer
    p_ft = footer.paragraphs[0]
    p_ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    txt = "지니매쓰 Premium" if is_commercial else "개인 학습용"
    if not is_commercial: set_read_only(doc)
    
    run_ft = p_ft.add_run(f"{txt}  |  ")
    set_font(run_ft, font_size=9, bold=True)
    add_page_number(run_ft)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


